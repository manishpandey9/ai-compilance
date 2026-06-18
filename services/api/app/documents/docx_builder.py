"""Build DOCX files from official section outlines."""

from __future__ import annotations

import io
from typing import Iterable

from docx import Document
from docx.shared import Pt

from app.documents.official_sections import OfficialSection
from app.documents.spec import RenderContext


def _add_placeholder(doc: Document, text: str = "[Complete with evidence.]") -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.italic = True


def _add_header_block(doc: Document, ctx: RenderContext, *, subtitle: str, legal_basis: str) -> None:
    doc.add_heading(subtitle, 0)
    doc.add_paragraph(f"Provider / organisation: {ctx.company_name}")
    doc.add_paragraph(f"AI system: {ctx.system_name}")
    doc.add_paragraph(f"Assessment reference: {ctx.assessment_id}")
    doc.add_paragraph(f"Risk tier: {ctx.risk_tier_label} | Role: {ctx.actor_role_label}")
    doc.add_paragraph(f"Legal basis: {legal_basis}")
    doc.add_paragraph(f"Source version: {ctx.source_version} | Rule set v{ctx.rule_version}")
    doc.add_paragraph(
        "Template aligned to Regulation (EU) 2024/1689. Complete all sections before conformity "
        "assessment or regulatory submission. Not legal advice."
    )


def render_official_sections_doc(
    ctx: RenderContext,
    *,
    title: str,
    legal_basis: str,
    sections: Iterable[OfficialSection],
) -> bytes:
    doc = Document()
    _add_header_block(doc, ctx, subtitle=title, legal_basis=legal_basis)
    for section in sections:
        doc.add_heading(f"{section.number}. {section.title}", level=1)
        p = doc.add_paragraph()
        run = p.add_run(section.legal_citation)
        run.italic = True
        run.font.size = Pt(9)
        for item in section.items:
            doc.add_paragraph(item, style="List Bullet")
        _add_placeholder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_numbered_list_doc(
    ctx: RenderContext,
    *,
    title: str,
    legal_basis: str,
    numbered_items: list[tuple[str, str]],
    extra_paragraphs: list[str] | None = None,
) -> bytes:
    doc = Document()
    _add_header_block(doc, ctx, subtitle=title, legal_basis=legal_basis)
    for num, text in numbered_items:
        doc.add_heading(f"{num}. {text}", level=2)
        _add_placeholder(doc)
    if extra_paragraphs:
        doc.add_heading("Notes", level=1)
        for para in extra_paragraphs:
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_bullet_checklist_doc(
    ctx: RenderContext,
    *,
    title: str,
    legal_basis: str,
    bullets: list[str],
) -> bytes:
    doc = Document()
    _add_header_block(doc, ctx, subtitle=title, legal_basis=legal_basis)
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")
        _add_placeholder(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
